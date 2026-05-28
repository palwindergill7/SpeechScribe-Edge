"""Render transcript text into PDF / DOCX / TXT / SRT files."""

import os
import re

OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def export_file(text: str, fmt: str, base_name: str) -> str:
    path = os.path.join(OUTPUT_DIR, f"{base_name}.{fmt}")
    if fmt == "pdf":
        _to_pdf(text, path)
    elif fmt == "docx":
        _to_docx(text, path)
    elif fmt == "txt":
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)
    elif fmt == "srt":
        _to_srt(text, path)
    else:
        raise ValueError(f"Unsupported format: {fmt}")
    return path


def _to_pdf(text: str, path: str) -> None:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet

    doc = SimpleDocTemplate(path)
    styles = getSampleStyleSheet()
    story = []
    for line in text.split("\n"):
        if line.strip() == "":
            story.append(Spacer(1, 10))
        elif line.lower().startswith("title"):
            story.append(Paragraph(f"<b>{line}</b>", styles["Title"]))
        elif ":" in line and len(line) < 80:
            story.append(Paragraph(f"<b>{line}</b>", styles["Heading2"]))
        else:
            story.append(Paragraph(line, styles["BodyText"]))
    doc.build(story)


def _to_docx(text: str, path: str) -> None:
    from docx import Document

    document = Document()
    for line in text.split("\n"):
        if line.strip() == "":
            document.add_paragraph("")
        elif line.lower().startswith("title"):
            document.add_heading(line, level=0)
        elif ":" in line and len(line) < 80:
            document.add_heading(line, level=2)
        else:
            document.add_paragraph(line)
    document.save(path)


def _to_srt(text: str, path: str) -> None:
    """
    If the transcript has [hh:mm:ss - hh:mm:ss] stamps, build real SRT cues.
    Otherwise fall back to one cue per non-empty line.
    """
    stamp_re = re.compile(r"\[(\d{1,2}:\d{2}(?::\d{2})?)\s*-\s*(\d{1,2}:\d{2}(?::\d{2})?)\]\s*(.*)")
    cues = []
    for line in text.split("\n"):
        if not line.strip():
            continue
        m = stamp_re.match(line.strip())
        if m:
            cues.append((_srt_time(m.group(1)), _srt_time(m.group(2)), m.group(3)))
        else:
            cues.append((None, None, line.strip()))

    with open(path, "w", encoding="utf-8") as f:
        for i, (start, end, content) in enumerate(cues, 1):
            if start is None:
                start = _srt_time_from_seconds((i - 1) * 3)
                end = _srt_time_from_seconds(i * 3)
            f.write(f"{i}\n{start} --> {end}\n{content}\n\n")


def _srt_time(stamp: str) -> str:
    parts = [int(p) for p in stamp.split(":")]
    if len(parts) == 2:
        m, s = parts
        h = 0
    else:
        h, m, s = parts
    return f"{h:02d}:{m:02d}:{s:02d},000"


def _srt_time_from_seconds(total: int) -> str:
    m, s = divmod(total, 60)
    h, m = divmod(m, 60)
    return f"{h:02d}:{m:02d}:{s:02d},000"
